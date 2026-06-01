from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "memoire_work"
OUT = WORK / "memoire_proxmox_sentinel_revise.docx"

FIGURES = {
    "architecture": Path(r"C:\Users\enzo2\Downloads\dsdf.png"),
    "cpu_ram": Path(
        r"C:\Users\enzo2\Downloads\figures_02-05-2026 (6)\analysis_output\poster_augmented\02-05-2026_augmented3_01_regles_cpu_ram_limites.png"
    ),
    "rules_matrix": Path(
        r"C:\Users\enzo2\Downloads\figures_02-05-2026 (6)\analysis_output\poster_augmented\02-05-2026_augmented3_02_regles_confusion_matrix.png"
    ),
    "iforest_score": Path(
        r"C:\Users\enzo2\Downloads\figures_02-05-2026 (6)\analysis_output\poster_augmented\02-05-2026_augmented3_04_iforest_score_anomalie.png"
    ),
    "iforest_matrix": Path(
        r"C:\Users\enzo2\Downloads\figures_02-05-2026 (6)\analysis_output\poster_augmented\02-05-2026_augmented3_05_iforest_confusion_matrix_projection.png"
    ),
    "delays": Path(
        r"C:\Users\enzo2\Downloads\figures_02-05-2026 (6)\analysis_output\poster_augmented\02-05-2026_augmented3_07_comparaison_delais_regles_vs_iforest.png"
    ),
}


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


def add_figure(doc: Document, key: str, caption: str, width: float = 6.1) -> None:
    fig = FIGURES[key]
    if not fig.exists():
        p = doc.add_paragraph()
        p.add_run(f"[Figure manquante : {fig}]").italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig), width=Inches(width))
    add_caption(doc, caption)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_small_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size, bold, small_caps in [
        ("Heading 1", 13, True, True),
        ("Heading 2", 12, True, False),
        ("Heading 3", 12, True, False),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.small_caps = small_caps
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    set_update_fields(doc)
    return doc


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run("Institut Catholique de Lille\nEDN - Master Cyber")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Développement d’un SOC automatisé sur Proxmox VE")
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Proxmox Sentinel : supervision, corrélation et réponse active human-in-the-loop")
    run.font.name = "Times New Roman"
    run.font.size = Pt(15)
    run.italic = True

    details = [
        "Mémoire de Master Cybersécurité",
        "Chevalier Enzo",
        "Année universitaire 2025-2026",
        "Version révisée - mai 2026",
    ]
    for item in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)

    add_small_note(
        doc,
        "Remarque : cette copie de travail conserve l’original intact et intègre les résultats expérimentaux du POC ainsi que les exigences principales du guide de livrable.",
    )
    doc.add_page_break()


def add_attestation(doc: Document) -> None:
    add_heading(doc, "Attestation de non-plagiat", 1)
    add_para(
        doc,
        "Je soussigné, Chevalier Enzo, certifie que ce mémoire est le résultat d’un travail personnel. Les sources utilisées, qu’elles soient scientifiques, techniques ou issues de la documentation logicielle, sont citées dans le corps du texte et référencées en bibliographie. Les figures, tableaux et résultats issus du prototype sont présentés comme tels, et les données augmentées ou prospectives sont explicitement distinguées des mesures brutes.",
    )
    add_para(
        doc,
        "L’usage d’outils d’assistance à la rédaction et à la programmation a servi à structurer, reformuler, vérifier et documenter le travail. La conception du POC, les campagnes de test, l’interprétation des résultats et la validation finale demeurent sous ma responsabilité.",
    )
    add_para(doc, "Fait à Lille, le 13 mai 2026.")
    add_para(doc, "Signature :")
    doc.add_page_break()


def add_summaries(doc: Document) -> None:
    add_heading(doc, "Résumé", 1)
    add_para(
        doc,
        "Ce mémoire étudie l’apport de l’automatisation pour accélérer la détection et la réaction face à des anomalies dans un environnement Proxmox VE. Le travail s’appuie sur un prototype nommé Proxmox Sentinel, développé en Python avec Streamlit, proxmoxer, Docker Compose et SQLite. Le système collecte en continu les métriques CPU/RAM de l’hyperviseur et des VM QEMU via l’API Proxmox, centralise les journaux SSH par Syslog, applique des règles de détection explicables et propose une réponse active contrôlée par l’analyste. La remédiation consiste à isoler ou restaurer l’interface réseau net0 d’une VM QEMU, en conservant une validation humaine afin de limiter les actions destructrices dans un contexte de preuve de concept.",
    )
    add_para(
        doc,
        "La méthodologie combine une baseline normale, des charges CPU légitimes, des scans Nmap, des tentatives Hydra, des scénarios mixtes et des tests de réponse active. Les résultats montrent que les métriques CPU/RAM seules détectent correctement les pressions de ressources, mais ne qualifient pas l’intention malveillante : Hydra peut rester invisible côté ressources tandis qu’une charge légitime déclenche une alerte. L’ajout des logs SSH via Syslog réduit ce faux négatif en rendant visibles les échecs d’authentification. Enfin, une extension Isolation Forest est évaluée sur des données historiques enrichies afin d’améliorer la priorisation des incidents. Le POC valide donc l’intérêt d’un SOC léger, traçable et adapté aux petits environnements virtualisés, tout en soulignant les limites des seuils simples et la nécessité d’enrichir progressivement les sources de données.",
    )
    add_para(doc, "Mots-clés : Proxmox VE ; SOC ; automatisation ; Syslog ; Isolation Forest.")

    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "This thesis investigates how automation can reduce anomaly detection and response delays in a Proxmox VE environment. The work is based on a proof of concept named Proxmox Sentinel, implemented in Python with Streamlit, proxmoxer, Docker Compose and SQLite. The system continuously collects CPU and memory metrics from the hypervisor and QEMU virtual machines through the Proxmox API, centralizes SSH authentication logs through Syslog, applies explainable detection rules and provides a human-in-the-loop active response workflow. The response action isolates or restores the net0 interface of a QEMU virtual machine while keeping analyst validation to avoid unsafe autonomous remediation in a laboratory proof of concept.",
    )
    add_para(
        doc,
        "The experimental methodology combines normal baselines, legitimate CPU loads, Nmap scans, Hydra brute-force attempts, mixed scenarios and active response tests. Results show that CPU/RAM metrics alone can detect resource pressure but cannot reliably infer malicious intent: a Hydra attack may remain invisible at the resource level, while a legitimate CPU load may trigger a false positive. Adding SSH logs through Syslog reduces this blind spot by exposing authentication failures. An Isolation Forest extension is then evaluated on enriched historical observations to improve incident prioritization. The prototype therefore demonstrates the value of a lightweight and traceable SOC for small virtualized environments, while highlighting the limits of simple thresholds and the need to progressively enrich data sources.",
    )
    add_para(doc, "Keywords: Proxmox VE; SOC; automation; Syslog; Isolation Forest.")
    doc.add_page_break()


def add_toc_page(doc: Document) -> None:
    add_heading(doc, "Sommaire", 1)
    p = doc.add_paragraph()
    add_toc(p)
    add_small_note(doc, "Dans Microsoft Word : clic droit sur le sommaire, puis « Mettre à jour les champs » pour afficher la pagination finale.")
    add_para(doc, "Plan de lecture :")
    for entry in [
        "1. Introduction",
        "2. Revue de la littérature et problématique",
        "3. Méthodologie",
        "4. Résultats",
        "5. Discussion",
        "6. Conclusion",
        "Bibliographie",
        "Annexes",
    ]:
        add_bullet(doc, entry)
    doc.add_page_break()


def add_introduction(doc: Document) -> None:
    add_heading(doc, "1. Introduction", 1)
    add_para(
        doc,
        "La virtualisation de type 1 s’est imposée dans les infrastructures professionnelles et dans les environnements auto-hébergés. Proxmox VE, fondé sur QEMU/KVM pour les machines virtuelles et LXC pour les conteneurs, offre une solution accessible pour consolider des services, créer des laboratoires et administrer plusieurs charges de travail depuis un même hyperviseur. Cette accessibilité a toutefois un coût opérationnel : les métriques, les journaux et les alertes peuvent rapidement dépasser ce qu’un administrateur seul peut surveiller manuellement.",
    )
    add_para(
        doc,
        "La question centrale de ce mémoire est donc la suivante : comment l’automatisation permet-elle d’accélérer la détection d’anomalies par rapport à une surveillance manuelle ? Cette question ne concerne pas seulement la détection, mais aussi la capacité à prioriser, tracer et déclencher une réponse active maîtrisée. Dans un SOC traditionnel, cette logique s’inscrit dans le prolongement des SIEM, XDR et SOAR, dont l’objectif est de réduire la charge cognitive, d’améliorer la coordination et de raccourcir les délais de réaction (Pitkar, 2025 ; Lee et al., 2022).",
    )
    add_para(
        doc,
        "Le travail présenté ici prend la forme d’un POC de mémoire, nommé Proxmox Sentinel. Il ne vise pas à remplacer une pile industrielle comme ELK, Prometheus ou Grafana, mais à démontrer qu’un SOC léger peut être conçu autour de quatre fonctions essentielles : collecte continue, détection explicable, journalisation des incidents et réponse active human-in-the-loop. Le périmètre retenu est volontairement restreint aux VM QEMU et à l’interface réseau net0 afin de préserver une démonstration réaliste, reproductible et défendable.",
    )
    add_para(
        doc,
        "L’objectif principal est de concevoir un outil capable de superviser un serveur Proxmox, de détecter des anomalies à partir de métriques et de logs, puis de proposer une action d’isolement réseau contrôlée. L’hypothèse étudiée est que l’automatisation, enrichie par la corrélation métriques + logs + ML, réduit le délai de détection et améliore la priorisation des incidents par rapport à une surveillance manuelle.",
    )


def add_literature(doc: Document) -> None:
    add_heading(doc, "2. Revue de la littérature et problématique", 1)
    add_heading(doc, "2.1 Supervision des environnements virtualisés", 2)
    add_para(
        doc,
        "Les environnements virtualisés imposent une difficulté particulière : une anomalie peut se manifester au niveau de l’hyperviseur, d’une VM ou d’un service invité. Les approches intrusives, comme l’introspection de VM, apportent une visibilité forte mais augmentent la complexité et les contraintes de déploiement. Mishra et al. (2020) illustrent cette famille de travaux avec KVMInspector, orienté détection de malware en environnement KVM. À l’inverse, les approches indirectes exploitent des métriques de ressources sans accès profond au système surveillé. Jindal et al. (2022) montrent ainsi qu’il est possible de détecter certaines anomalies VMM à partir de signaux indirects observés sur les VM, ce qui correspond bien à un POC agentless pour Proxmox.",
    )
    add_para(
        doc,
        "Ce mémoire s’inscrit dans cette seconde logique : le prototype interroge l’API Proxmox pour récupérer les états de l’hôte et des VM, sans installer d’agent de monitoring lourd dans la cible. Cette posture réduit la friction de déploiement, mais limite aussi la profondeur d’analyse. Une anomalie réseau peu coûteuse, comme un scan ou un brute-force SSH de faible intensité, ne provoque pas nécessairement de pic CPU/RAM détectable.",
    )
    add_heading(doc, "2.2 Automatisation SOC, triage et surcharge analyste", 2)
    add_para(
        doc,
        "L’automatisation sécurité est généralement justifiée par la réduction du temps de détection, du temps de réaction et de la charge de triage. Les travaux sur les architectures SIEM/SOAR insistent sur la coordination entre collecte, corrélation et réponse (Pitkar, 2025 ; Lee et al., 2022). Cette coordination ne signifie pas nécessairement une remédiation totalement autonome. Dans un contexte académique et de laboratoire, un modèle human-in-the-loop reste plus prudent : le système met en évidence l’incident, documente le contexte et propose une action, mais l’analyste valide l’isolement.",
    )
    add_para(
        doc,
        "La priorisation est un enjeu central. Lee et al. (2022), avec SIERRA, soulignent que les analystes ne peuvent pas traiter tous les événements et qu’un système utile doit classer les anomalies selon leur intérêt opérationnel. De même, Liu et al. (2022) abordent l’accélération du triage par apprentissage de contexte. Dans le POC, cette idée se traduit par des scores de sévérité, des statuts d’incident et un poste de traitement qui concentre les alertes, la timeline et les actions disponibles.",
    )
    add_heading(doc, "2.3 Détection d’anomalies, règles et apprentissage automatique", 2)
    add_para(
        doc,
        "La détection d’intrusion par anomalie repose sur une comparaison entre un comportement attendu et un comportement observé. Les synthèses de Chandola et al. (2009), Buczak et Guven (2016) et García-Teodoro et al. (2009) rappellent que cette famille d’approches est pertinente pour découvrir des comportements inattendus, mais qu’elle expose aussi au risque de faux positifs. Axelsson (2000) montre que ce problème est structurel dans les systèmes de détection, notamment lorsque la base d’événements malveillants réels est faible.",
    )
    add_para(
        doc,
        "Les seuils explicables restent utiles pour une première version : ils sont simples à auditer et adaptés à une soutenance. Leur limite est qu’ils détectent une pression de ressource sans comprendre l’intention. L’apprentissage automatique peut enrichir cette logique. Isolation Forest (Liu et al., 2008) est particulièrement adapté aux données non étiquetées, car il isole les observations rares sans nécessiter une base complète d’attaques. Random Forest (Breiman, 2001) est également pertinent lorsque des données étiquetées suffisantes existent. Iacovazzi et Raza (2022) combinent ces deux familles pour la détection d’intrusion en conteneurs à partir de comportements système, ce qui justifie l’extension ML envisagée dans ce mémoire.",
    )
    add_para(
        doc,
        "Enfin, les attaques SSH constituent un cas pratique de corrélation. Javed et Paxson (2013), Park et al. (2021) et Hubballi et al. (2020) montrent que les brute-force distribués ou discrets nécessitent une analyse des échecs d’authentification, des sources et des fenêtres temporelles. Cette littérature explique pourquoi l’ajout de Syslog dans le POC est décisif : il corrige une limite observée avec Hydra, qui peut rester discret côté CPU/RAM.",
    )
    add_heading(doc, "2.4 Problématique et hypothèse", 2)
    add_para(
        doc,
        "La problématique retenue est : comment l’automatisation permet-elle d’accélérer la détection d’anomalies par rapport à une surveillance manuelle ? L’hypothèse testée est double. Premièrement, une collecte continue et des règles simples réduisent le délai de détection des anomalies de ressources. Deuxièmement, la corrélation des métriques Proxmox avec les logs SSH, puis avec un score Isolation Forest, améliore la priorisation des incidents et réduit les faux positifs ou faux négatifs les plus visibles dans le POC.",
    )


def add_methodology(doc: Document) -> None:
    add_heading(doc, "3. Méthodologie", 1)
    add_heading(doc, "3.1 Environnement expérimental", 2)
    add_para(
        doc,
        "L’environnement de test repose sur un serveur Proxmox VE de laboratoire hébergeant au minimum une VM cible QEMU Linux et une VM dédiée au SOC Dashboard. L’attaquant est simulé depuis un poste externe au serveur, via Exegol et des outils de type Nmap et Hydra. Cette séparation permet de distinguer le périmètre surveillé, la source de trafic offensif et le composant d’analyse. Les tests ne ciblent que la VM de démonstration et non l’hôte Proxmox.",
    )
    add_para(
        doc,
        "La VM SOC exécute Docker Compose avec trois services principaux : l’interface Streamlit, le collecteur Proxmox et le collecteur Syslog. La configuration est centralisée dans un fichier .env. L’accès à Proxmox se fait par un utilisateur dédié soc-dashboard@pve et un token API. Le paramètre VERIFY_SSL=False est accepté uniquement dans le contexte de lab et doit être remplacé par une validation TLS correcte dans un environnement de production.",
    )
    add_figure(
        doc,
        "architecture",
        "Figure 1 - Architecture technique du POC Proxmox Sentinel.",
        width=6.5,
    )

    add_heading(doc, "3.2 Architecture logicielle", 2)
    add_para(
        doc,
        "Le code a été organisé pour séparer l’interface, la collecte, la détection, le stockage et les actions de réponse. Cette séparation permet de faire tourner la collecte même lorsque la page Streamlit n’est pas ouverte. Streamlit devient alors une interface de supervision et de traitement incident, tandis que les services de fond alimentent la base SQLite.",
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_cell_text(hdr.cells[0], "Module", True)
    set_cell_text(hdr.cells[1], "Rôle", True)
    set_cell_text(hdr.cells[2], "Apport pour le mémoire", True)
    rows = [
        ("proxmox_client.py", "Connexion à l’API Proxmox et lecture des métriques.", "Collecte agentless, cohérente avec l’approche indirecte."),
        ("collector.py", "Boucle de collecte continue des noeuds et VM QEMU.", "Mesure indépendante de l’ouverture de l’interface."),
        ("syslog_collector.py", "Réception TCP/UDP des logs SSH via rsyslog.", "Corrélation métriques + logs pour réduire les faux négatifs."),
        ("detection.py", "Règles de seuil, score et corrélation SSH/CPU.", "Détection explicable avant l’extension ML."),
        ("ml_model.py", "Score Isolation Forest sur observations enrichies.", "Perspective ML pour améliorer la priorisation."),
        ("storage.py", "Persistance SQLite des métriques, alertes, incidents et actions.", "Traçabilité, chronologie et calculs MTTD/MTTR."),
        ("actions.py", "Isolement/restauration réseau via net0.", "Réponse active contrôlée et reproductible."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    add_caption(doc, "Tableau 1 - Organisation modulaire du prototype.")

    add_heading(doc, "3.3 Données collectées et persistance", 2)
    add_para(
        doc,
        "Les données persistées dans SQLite comprennent les métriques hôte et VM, les alertes générées, les incidents, les événements SSH, les actions d’isolement/restauration, les statuts de collecteur et les scores ML. Le choix de SQLite est volontairement léger : il suffit au POC, facilite la sauvegarde de la base et évite l’introduction d’une pile lourde qui masquerait l’objet de recherche.",
    )
    add_para(
        doc,
        "Les logs SSH sont reçus par Syslog sur le port 5514. Chaque VM Linux surveillée peut transmettre ses événements auth/authpriv via rsyslog. Le parser extrait les événements de type failed_password, invalid_user et accepted_password, puis les associe à une VM à l’aide d’une cartographie manuelle IP/VMID. Cette solution évite les connexions SSH répétées depuis le SOC vers les VM et rend la collecte plus pérenne.",
    )

    add_heading(doc, "3.4 Moteur de détection et réponse active", 2)
    add_para(
        doc,
        "Le moteur de détection utilise d’abord des règles configurables : seuil CPU hôte, seuil CPU VM, seuil RAM VM, durée minimale avant alerte, échecs SSH par fenêtre temporelle et corrélation SSH + CPU. Les alertes sont classées selon une sévérité low, medium ou critical, avec un score explicable. Cette couche par règles sert de base interprétable avant toute extension ML.",
    )
    add_para(
        doc,
        "La réponse active est limitée aux VM QEMU et à net0. L’action d’isolement modifie la configuration Proxmox de l’interface réseau en ajoutant link_down=1 ; la restauration supprime ce paramètre. Les VM critiques peuvent être exclues via PROTECTED_VMIDS. Le workflow impose une confirmation humaine : prendre en charge l’incident, décider de l’isolement, vérifier l’état réseau, restaurer, puis clore l’incident.",
    )

    add_heading(doc, "3.5 Protocole expérimental", 2)
    add_para(
        doc,
        "Le protocole de test a été construit pour comparer des périodes normales, des charges légitimes et des comportements malveillants simulés. Les campagnes comprennent : baseline normale, maintenance légère, charge CPU contrôlée avec la commande yes, scans Nmap, brute-force SSH avec Hydra, scénarios mixtes Hydra + CPU et tests d’isolement/restauration. Les heures de début et de fin ont été consignées dans un fichier experiment_log.csv afin d’associer les fenêtres expérimentales aux données SQLite.",
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(["Scénario", "Objectif", "Résultat attendu", "Interprétation"]):
        set_cell_text(hdr.cells[i], text, True)
    rows = [
        ("Baseline", "Mesurer le bruit normal.", "Pas ou peu d’alertes.", "Vérifier le taux de faux positifs au repos."),
        ("Charge CPU légitime", "Tester les seuils ressources.", "Alertes CPU.", "Identifier les faux positifs liés à l’intention."),
        ("Nmap/Hydra seuls", "Simuler des attaques réseau peu coûteuses.", "Faible impact CPU/RAM.", "Mesurer les faux négatifs des métriques seules."),
        ("Hydra + Syslog", "Valider la corrélation logs SSH.", "Alerte ssh_intrusion.", "Réduire les faux négatifs observés."),
        ("Hydra + CPU + réponse", "Tester priorisation et isolement.", "Incident et action net0.", "Mesurer MTTD/MTTR et traçabilité."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    add_caption(doc, "Tableau 2 - Protocole expérimental synthétique.")


def add_results(doc: Document) -> None:
    add_heading(doc, "4. Résultats", 1)
    add_heading(doc, "4.1 Collecte continue et baseline", 2)
    add_para(
        doc,
        "La première validation concerne la collecte continue. Après ajout d’un service proxmox-collector indépendant, la base SQLite continue de recevoir des métriques même lorsque la page Streamlit n’est pas ouverte. Une baseline nocturne a produit 63 432 métriques sans alerte. Ce résultat montre que la collecte n’est plus dépendante du navigateur et qu’elle peut alimenter une preuve expérimentale exploitable.",
    )
    add_para(
        doc,
        "Les périodes de maintenance légère, comme la connexion SSH normale, apt update, l’installation de htop et des commandes système simples, n’ont pas généré d’alerte inattendue. Cette observation confirme que les seuils ne déclenchent pas sur un repos ou un usage très modéré du lab. Elle ne suffit pas à prouver l’absence de faux positifs en production, mais constitue une première baseline contrôlée.",
    )

    add_heading(doc, "4.2 Limite des métriques CPU/RAM seules", 2)
    add_para(
        doc,
        "Les cinq essais de charge CPU légitime ont généré cinq alertes, dont quatre critical et une medium résolues. Ces alertes prouvent que le SOC détecte bien une pression de ressource. Leur classification expérimentale est toutefois celle de faux positifs, car l’activité était volontaire et non malveillante. À l’inverse, les scans Nmap et certains essais Hydra seuls n’ont pas déclenché d’alerte côté CPU/RAM. La figure 2 illustre cette limite : Hydra seul reste discret sur les métriques, tandis qu’une charge CPU légitime franchit les seuils.",
    )
    add_figure(
        doc,
        "cpu_ram",
        "Figure 2 - Les métriques CPU/RAM détectent la pression ressource mais ne qualifient pas l’intention.",
        width=6.3,
    )

    add_heading(doc, "4.3 Apport de Syslog SSH", 2)
    add_para(
        doc,
        "L’ajout du collecteur Syslog transforme le comportement observé. Un test logger a validé la remontée d’un événement failed_password. Six échecs SSH manuels ont ensuite déclenché une alerte ssh_bruteforce_suspected, résolue après disparition des conditions. Les scénarios Hydra courts ont également créé des incidents ssh_intrusion. Cette évolution montre que les logs d’authentification rendent visibles des attaques peu coûteuses côté ressources, ce qui répond directement à la limite observée avec les métriques seules.",
    )
    add_para(
        doc,
        "Dans le scénario Hydra + CPU, le SOC a généré simultanément des alertes liées aux échecs SSH, à la pression CPU et à la corrélation ssh_cpu_correlated. Cette corrélation augmente la sévérité et donne un meilleur contexte à l’analyste. Elle correspond à l’idée de priorisation défendue dans les travaux sur SIEM/SOAR et sur le ranking d’activités anormales.",
    )
    add_figure(
        doc,
        "rules_matrix",
        "Figure 3 - Matrice expérimentale extrapolée pour la détection par règles sur 100 scénarios.",
        width=5.1,
    )
    add_small_note(
        doc,
        "Les matrices sur 100 scénarios sont des extrapolations équilibrées construites à partir des observations de campagne. Elles servent à comparer les tendances, pas à prétendre à une mesure statistique exhaustive.",
    )

    add_heading(doc, "4.4 Extension Isolation Forest", 2)
    add_para(
        doc,
        "Une extension Isolation Forest a été introduite afin de scorer les comportements à partir des métriques et des logs SSH. En baseline, les scores live observés sur les VM restaient faibles, autour de 12 à 15, ce qui correspond à un comportement normal. Lors d’une attaque Hydra, le score oscillait entre 65 et 75 avec le seuil initial, montrant une détection sensible mais instable. Un seuil de 60 a ensuite été utilisé pour illustrer une configuration plus adaptée au jeu de données augmenté.",
    )
    add_para(
        doc,
        "L’intérêt de cette approche est de ne pas dépendre uniquement d’un seuil CPU. Le score agrège la forme du comportement : échecs SSH, activité ressource, contexte temporel et retour au calme. La figure 4 montre que le score signale Hydra même sans pic CPU, puis priorise Hydra + CPU. La figure 5 compare la matrice obtenue avec l’approche par règles : sur l’extrapolation à 100 scénarios, le nombre de faux positifs diminue de 22 à 4, avec une précision passant de 66 % à 91 %, tout en conservant un rappel de 84 %.",
    )
    add_figure(
        doc,
        "iforest_score",
        "Figure 4 - Score d’anomalie Isolation Forest sur les fenêtres expérimentales augmentées.",
        width=6.3,
    )
    add_figure(
        doc,
        "iforest_matrix",
        "Figure 5 - Matrice prospective Isolation Forest sur 100 scénarios augmentés.",
        width=5.1,
    )

    add_heading(doc, "4.5 Délais de détection et de réaction", 2)
    add_para(
        doc,
        "La comparaison MTTD/MTTR synthétise l’apport progressif de l’automatisation. Une surveillance manuelle est représentée par un délai de détection plus long, car l’analyste doit observer les métriques et les logs sans corrélation automatique. Le SOC par règles + Syslog réduit fortement le MTTD, car l’incident est créé dès que les conditions de détection sont réunies. L’extension Isolation Forest améliore encore la priorisation, ce qui peut réduire le temps nécessaire pour identifier l’incident à traiter.",
    )
    add_figure(
        doc,
        "delays",
        "Figure 6 - Comparaison MTTD/MTTR entre surveillance manuelle, règles + Syslog et Isolation Forest.",
        width=6.3,
    )
    add_para(
        doc,
        "La réponse active a été validée sur une VM QEMU de démonstration. L’isolement de net0 a coupé le réseau de la VM cible, puis la restauration a rétabli la connectivité. Les actions sont journalisées dans SQLite avec horodatage, résultat et VMID. Le POC ne ferme pas automatiquement l’incident : l’analyste confirme la prise en charge, le confinement, la restauration et la clôture. Ce choix rend la démonstration plus sûre et plus défendable.",
    )


def add_discussion(doc: Document) -> None:
    add_heading(doc, "5. Discussion", 1)
    add_heading(doc, "5.1 Réponse à la problématique", 2)
    add_para(
        doc,
        "Les résultats confirment que l’automatisation accélère la détection d’anomalies par rapport à une surveillance manuelle, à condition que les sources soient adaptées à l’anomalie recherchée. Les métriques CPU/RAM suffisent pour détecter une pression de ressource, mais pas pour inférer l’intention. L’ajout des logs SSH réduit les faux négatifs sur les attaques Hydra. L’Isolation Forest, utilisée comme extension de priorisation, réduit surtout les faux positifs dans les scénarios augmentés.",
    )
    add_para(
        doc,
        "Le POC rejoint les constats de la littérature : les systèmes de détection par anomalie sont utiles mais sensibles aux faux positifs (Axelsson, 2000 ; Chandola et al., 2009), l’apprentissage non supervisé est pertinent lorsque les données étiquetées sont rares (Liu et al., 2008), et la priorisation est essentielle pour réduire la surcharge analyste (Lee et al., 2022). Dans le contexte Proxmox, l’approche indirecte par métriques reste pertinente pour un déploiement léger, comme le suggèrent les travaux de Jindal et al. (2022), mais elle doit être enrichie par des logs de sécurité.",
    )
    add_heading(doc, "5.2 Apports du prototype", 2)
    add_para(
        doc,
        "Le premier apport est technique : le prototype montre qu’un SOC léger peut être construit avec une architecture modulaire, conteneurisée et persistante. Le second apport est méthodologique : la base SQLite et le fichier de journal expérimental permettent de produire des preuves, des timelines, des matrices et des indicateurs MTTD/MTTR. Le troisième apport est opérationnel : le poste incident rapproche détection, analyse, isolement et restauration dans un même parcours utilisateur.",
    )
    add_para(
        doc,
        "Le choix human-in-the-loop est également un apport. Dans un lab, il serait possible d’isoler automatiquement une VM dès l’apparition d’une alerte critique. Cependant, une remédiation autonome mal calibrée peut interrompre un service légitime. Le POC privilégie donc une automatisation d’aide à la décision : elle accélère l’identification et prépare l’action, mais conserve une validation humaine pour les décisions impactantes.",
    )
    add_heading(doc, "5.3 Limites", 2)
    add_para(
        doc,
        "Plusieurs limites doivent être soulignées. Le périmètre de réponse est limité aux VM QEMU et à net0. Les LXC ne sont pas traités. Le POC ne réalise pas d’inspection réseau profonde, ne collecte pas encore tous les journaux système et ne gère pas Syslog chiffré. Les scénarios sont contrôlés et limités à un lab. Les matrices à 100 scénarios sont des données augmentées à partir des observations ; elles servent à illustrer une tendance et devront être validées par des campagnes plus larges.",
    )
    add_para(
        doc,
        "L’Isolation Forest améliore la priorisation, mais son intégration reste une première extension. Le modèle doit être réentraîné sur des baselines plus longues et plus variées pour limiter les dérives. D’autres modèles, comme Random Forest avec données étiquetées ou des approches séquentielles, pourraient être comparés lorsque le volume de données sera suffisant.",
    )


def add_conclusion(doc: Document) -> None:
    add_heading(doc, "6. Conclusion", 1)
    add_para(
        doc,
        "Ce mémoire a présenté la conception et l’évaluation d’un POC de SOC automatisé pour Proxmox VE. Le prototype valide l’intérêt d’une collecte continue, d’une corrélation métriques + logs, d’une priorisation par score et d’une réponse active contrôlée. Les campagnes montrent que l’automatisation réduit le MTTD, structure le traitement incident et rend les décisions plus traçables.",
    )
    add_para(
        doc,
        "La réponse à la problématique est donc nuancée : l’automatisation accélère bien la détection, mais seulement si les signaux collectés couvrent le comportement étudié. Les métriques ressources seules sont insuffisantes pour détecter des attaques SSH discrètes. Syslog corrige cette limite, et l’Isolation Forest apporte une piste pour mieux prioriser les incidents et réduire les faux positifs. Les perspectives portent sur l’amélioration du modèle ML, l’ajout de logs réseau et système, la mise en place de playbooks plus fins, l’action sur l’IP attaquante plutôt que sur la VM cible, et l’extension progressive vers une remédiation plus autonome après validation expérimentale.",
    )


def add_bibliography(doc: Document) -> None:
    add_heading(doc, "Bibliographie", 1)
    refs = [
        ("Axelsson, S. (2000). The base-rate fallacy and the difficulty of intrusion detection. ACM Transactions on Information and System Security, 3(3), 186-205.", "https://doi.org/10.1145/357830.357849"),
        ("Breiman, L. (2001). Random forests. Machine Learning, 45, 5-32.", "https://doi.org/10.1023/A:1010933404324"),
        ("Buczak, A. L., & Guven, E. (2016). A survey of data mining and machine learning methods for cyber security intrusion detection. IEEE Communications Surveys & Tutorials, 18(2), 1153-1176.", "https://doi.org/10.1109/COMST.2015.2494502"),
        ("Chandola, V., Banerjee, A., & Kumar, V. (2009). Anomaly detection: A survey. ACM Computing Surveys, 41(3), Article 15.", "https://doi.org/10.1145/1541880.1541882"),
        ("García-Teodoro, P., Díaz-Verdejo, J., Maciá-Fernández, G., & Vázquez, E. (2009). Anomaly-based network intrusion detection: Techniques, systems and challenges. Computers & Security, 28(1-2), 18-28.", "https://doi.org/10.1016/j.cose.2008.08.003"),
        ("Hubballi, N., Tiwari, N., & Khandait, P. (2020). Distributed SSH bruteforce attack detection with flow content similarity and login failure reputation. ASIA CCS 2020, 916-918.", "https://doi.org/10.1145/3320269.3405443"),
        ("Iacovazzi, A., & Raza, S. (2022). Ensemble of random and isolation forests for graph-based intrusion detection in containers. 2022 IEEE International Conference on Cyber Security and Resilience, 30-37.", "https://doi.org/10.1109/CSR54599.2022.9850307"),
        ("Javed, M., & Paxson, V. (2013). Detecting stealthy, distributed SSH brute-forcing. Proceedings of the 2013 ACM SIGSAC Conference on Computer & Communications Security, 85-96.", "https://doi.org/10.1145/2508859.2516719"),
        ("Jindal, A., Shakhat, I., Cardoso, J., Gerndt, M., & Podolskiy, V. (2022). IAD: Indirect anomalous VMMs detection in the cloud-based environment. Service-Oriented Computing - ICSOC 2021 Workshops, 190-201.", "https://doi.org/10.1007/978-3-031-14135-5_15"),
        ("Lee, J., Tang, F., Thet, P. M., Yeoh, D., Rybczynski, M., & Divakaran, D. M. (2022). SIERRA: Ranking anomalous activities in enterprise networks. IEEE European Symposium on Security and Privacy.", "https://doi.org/10.1109/eurosp53844.2022.00011"),
        ("Lee, M., Jang-Jaccard, J., & Kwak, J. (2022). Novel architecture of security orchestration, automation and response in Internet of Blended Environment. Computers, Materials & Continua, 73(1), 199-223.", "https://doi.org/10.32604/cmc.2022.028495"),
        ("Liu, F. T., Ting, K. M., & Zhou, Z.-H. (2008). Isolation Forest. 2008 IEEE International Conference on Data Mining, 413-422.", "https://doi.org/10.1109/ICDM.2008.17"),
        ("Liu, J., Zhang, R., Liu, W., Zhang, Y., Yang, H., & Ren, J. (2022). Context2Vector: Accelerating security event triage via context representation learning. Information and Software Technology, 146, 106856.", "https://doi.org/10.1016/j.infsof.2022.106856"),
        ("Milenkoski, A., Vieira, M., Kounev, S., Avritzer, A., & Payne, B. D. (2015). Evaluating computer intrusion detection systems: A survey of common practices. ACM Computing Surveys, 48(1), Article 12.", "https://doi.org/10.1145/2808691"),
        ("Mishra, P., Verma, I., & Gupta, S. (2020). KVMInspector: KVM based introspection approach to detect malware in cloud environment. Journal of Information Security and Applications, 51, 102460.", "https://doi.org/10.1016/j.jisa.2020.102460"),
        ("Ntambu, P., & Adeshina, S. A. (2021). Machine learning-based anomalies detection in cloud virtual machine resource usage. 2021 International Conference on Mechanical Engineering, Automation and Systems.", "https://doi.org/10.1109/ICMEAS52683.2021.9692308"),
        ("Park, K., Song, Y., & Cheong, Y.-G. (2021). Network log-based SSH brute-force attack detection model. Computers, Materials & Continua, 68(1), 887-901.", "https://doi.org/10.32604/cmc.2021.015172"),
        ("Pitkar, H. (2025). Cloud security automation through symmetry: Threat detection and response. Symmetry, 17(6), 859.", "https://doi.org/10.3390/sym17060859"),
        ("Sommer, R., & Paxson, V. (2010). Outside the closed world: On using machine learning for network intrusion detection. IEEE Symposium on Security and Privacy, 305-316.", "https://doi.org/10.1109/SP.2010.25"),
    ]
    for ref, doi in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"DOI : {doi}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)


def add_annexes(doc: Document) -> None:
    add_heading(doc, "Annexes", 1)
    add_heading(doc, "Annexe A - Campagnes de test", 2)
    add_para(
        doc,
        "Les principales campagnes de test ont été réalisées entre le 25 avril et le 2 mai 2026. Elles comprennent une baseline initiale, une baseline nocturne, cinq charges CPU légitimes, des scans Nmap, des brute-force SSH Hydra, des scénarios mixtes et des tests d’isolement/restauration. Les fenêtres ont été consignées dans experiment_log.csv afin de relier les événements aux métriques et alertes de la base SQLite.",
    )
    add_bullet(doc, "25-26 avril : baseline et charges CPU légitimes. Résultat : 63 432 métriques sur baseline nocturne, 0 alerte ; puis 5 alertes sur 5 charges CPU.")
    add_bullet(doc, "29 avril : Nmap/Hydra depuis Exegol. Résultat : attaques réseau peu visibles côté CPU/RAM.")
    add_bullet(doc, "1er mai : validation Syslog SSH, échecs manuels, Hydra, Hydra + CPU et isolement/restauration.")
    add_bullet(doc, "2 mai : validation du parcours incident UX, Hydra court, Hydra + CPU + isolement.")

    add_heading(doc, "Annexe B - Paramètres techniques principaux", 2)
    add_bullet(doc, "Authentification Proxmox : utilisateur soc-dashboard@pve et token API.")
    add_bullet(doc, "Collecte Proxmox : proxmox-collector en service Docker indépendant.")
    add_bullet(doc, "Collecte logs : soc-syslog sur 5514 TCP/UDP, alimenté par rsyslog côté VM.")
    add_bullet(doc, "Réponse active : VM QEMU uniquement, interface net0 uniquement, validation humaine obligatoire.")
    add_bullet(doc, "Persistance : SQLite avec métriques, alertes, incidents, événements SSH, actions, scores ML.")

    add_heading(doc, "Annexe C - Résumé des modifications apportées", 2)
    add_para(
        doc,
        "La version révisée remplace les formulations obsolètes liées aux LXC/Kali/InfluxDB par le périmètre réellement implémenté : VM QEMU, Exegol, SQLite, Syslog et Docker Compose. La problématique, la méthodologie, les résultats et la discussion ont été réalignés sur les campagnes de test. La bibliographie a été reconstruite avec au moins quinze références scientifiques et un DOI pour chaque article.",
    )


def build() -> None:
    doc = setup_document()
    add_cover(doc)
    add_attestation(doc)
    add_summaries(doc)
    add_toc_page(doc)
    add_introduction(doc)
    add_literature(doc)
    add_methodology(doc)
    add_results(doc)
    add_discussion(doc)
    add_conclusion(doc)
    doc.add_page_break()
    add_bibliography(doc)
    doc.add_page_break()
    add_annexes(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
